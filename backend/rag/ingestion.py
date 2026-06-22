"""Document extraction, chunking, storage, and metadata persistence."""

from __future__ import annotations

import csv
import io
import re
from dataclasses import dataclass
from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy.orm import Session

from backend.models.document import Document, DocumentChunk

SUPPORTED_FORMATS = frozenset({"pdf", "docx", "txt", "csv"})


class IngestionError(ValueError):
    """Base exception for invalid document ingestion requests."""


class UnsupportedDocumentError(IngestionError):
    """Raised when a document format is not supported."""


class DocumentExtractionError(IngestionError):
    """Raised when text cannot be extracted from a document."""


class EmptyDocumentError(IngestionError):
    """Raised when a document contains no extractable text."""


@dataclass(frozen=True)
class IngestionResult:
    """Identifier and chunk count for an ingested document."""

    document_id: str
    chunk_count: int


class DocumentIngestionService:
    """Coordinates document extraction, chunking, and persistence."""

    def __init__(
        self,
        upload_directory: Path,
        chunk_size: int = 1_000,
        chunk_overlap: int = 150,
    ) -> None:
        if chunk_size <= 0:
            raise ValueError("chunk_size must be greater than zero")
        if chunk_overlap < 0 or chunk_overlap >= chunk_size:
            raise ValueError("chunk_overlap must be between zero and chunk_size")

        self.upload_directory = upload_directory
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def ingest(
        self,
        *,
        filename: str,
        content: bytes,
        content_type: str | None,
        session: Session,
    ) -> IngestionResult:
        """Extract, chunk, save, and persist a document atomically."""
        safe_filename = Path(filename).name
        extension = Path(safe_filename).suffix.lower().lstrip(".")
        if extension not in SUPPORTED_FORMATS:
            raise UnsupportedDocumentError(
                f"Unsupported document format: {extension or 'missing extension'}"
            )

        text = self.extract_text(content, extension)
        chunks = self.chunk_text(text)
        if not chunks:
            raise EmptyDocumentError("The document contains no extractable text")

        document_id = str(uuid4())
        stored_filename = f"{document_id}.{extension}"
        self.upload_directory.mkdir(parents=True, exist_ok=True)
        file_path = self.upload_directory / stored_filename

        try:
            file_path.write_bytes(content)
            document = Document(
                id=document_id,
                original_filename=safe_filename,
                stored_filename=stored_filename,
                file_format=extension,
                content_type=content_type,
                size_bytes=len(content),
                file_path=str(file_path),
                chunk_count=len(chunks),
            )
            document.chunks = [
                DocumentChunk(
                    chunk_index=index,
                    text=chunk,
                    character_count=len(chunk),
                )
                for index, chunk in enumerate(chunks)
            ]
            session.add(document)
            session.commit()
        except Exception:
            session.rollback()
            file_path.unlink(missing_ok=True)
            raise

        return IngestionResult(document_id=document_id, chunk_count=len(chunks))

    def extract_text(self, content: bytes, file_format: str) -> str:
        """Extract normalized text from a supported document byte stream."""
        try:
            if file_format == "pdf":
                return self._extract_pdf(content)
            if file_format == "docx":
                return self._extract_docx(content)
            if file_format == "txt":
                return self._decode_text(content)
            if file_format == "csv":
                return self._extract_csv(content)
        except IngestionError:
            raise
        except Exception as exc:
            raise DocumentExtractionError(
                f"Could not extract text from {file_format.upper()} document"
            ) from exc

        raise UnsupportedDocumentError(f"Unsupported document format: {file_format}")

    def chunk_text(self, text: str) -> list[str]:
        """Split normalized text into overlapping, boundary-aware chunks."""
        normalized = re.sub(r"[ \t]+", " ", text)
        normalized = re.sub(r"\n{3,}", "\n\n", normalized).strip()
        if not normalized:
            return []

        chunks: list[str] = []
        start = 0
        text_length = len(normalized)

        while start < text_length:
            proposed_end = min(start + self.chunk_size, text_length)
            end = proposed_end

            if proposed_end < text_length:
                search_start = start + max(self.chunk_size // 2, 1)
                boundary = max(
                    normalized.rfind("\n\n", search_start, proposed_end),
                    normalized.rfind(". ", search_start, proposed_end),
                    normalized.rfind(" ", search_start, proposed_end),
                )
                if boundary > start:
                    end = boundary + (1 if normalized[boundary] == "." else 0)

            chunk = normalized[start:end].strip()
            if chunk:
                chunks.append(chunk)
            if end >= text_length:
                break

            next_start = max(end - self.chunk_overlap, start + 1)
            while next_start < end and normalized[next_start].isspace():
                next_start += 1
            start = next_start

        return chunks

    @staticmethod
    def _extract_pdf(content: bytes) -> str:
        reader = PdfReader(io.BytesIO(content))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)

    @staticmethod
    def _extract_docx(content: bytes) -> str:
        document = DocxDocument(io.BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_rows = [
            "\t".join(cell.text for cell in row.cells)
            for table in document.tables
            for row in table.rows
        ]
        return "\n".join(paragraphs + table_rows)

    def _extract_csv(self, content: bytes) -> str:
        decoded = self._decode_text(content)
        rows = csv.reader(io.StringIO(decoded))
        return "\n".join("\t".join(cell.strip() for cell in row) for row in rows)

    @staticmethod
    def _decode_text(content: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-16"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise DocumentExtractionError("Text file encoding must be UTF-8 or UTF-16")

